# api.py  v2.0 — PathForge backend
# Run: python -m uvicorn api:app --reload --port 8000

import os, sys, re, json, base64, traceback, io
from pathlib import Path
from fastapi import FastAPI, Form
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

app = FastAPI(title="PathForge Engine v2", version="2.0")
app.add_middleware(CORSMiddleware,
    allow_origins=["http://localhost:5173","http://localhost:3000","http://127.0.0.1:5173"],
    allow_credentials=True, allow_methods=["*"], allow_headers=["*"])

# ══════════════════════════════════════════════════════════════════════════════
# SKILL TAXONOMY  —  canonical name → (category, [domain_tags])
# ══════════════════════════════════════════════════════════════════════════════
SKILL_TAXONOMY = {
    "Python":("Programming",["tech","ml","data"]),
    "SQL":("Database",["tech","data","finance"]),
    "JavaScript":("Programming",["tech","web"]),
    "TypeScript":("Programming",["tech","web"]),
    "Java":("Programming",["tech"]),
    "C++":("Programming",["tech"]),
    "C#":("Programming",["tech"]),
    "Go":("Programming",["tech"]),
    "Rust":("Programming",["tech"]),
    "Scala":("Programming",["tech","data"]),
    "R":("Programming",["data","statistics"]),
    "MATLAB":("Programming",["science"]),
    "VBA":("Programming",["finance"]),
    "Bash":("Tools",["tech"]),
    "PHP":("Programming",["web"]),
    "Ruby":("Programming",["web"]),
    "Swift":("Programming",["mobile"]),
    "Kotlin":("Programming",["mobile"]),
    "HTML":("Frontend",["web"]),
    "CSS":("Frontend",["web"]),
    "React":("Frontend",["web"]),
    "Vue":("Frontend",["web"]),
    "Angular":("Frontend",["web"]),
    "Next.js":("Frontend",["web"]),
    "Node.js":("Backend",["web"]),
    "jQuery":("Frontend",["web"]),
    "Bootstrap":("Frontend",["web"]),
    "Tailwind":("Frontend",["web"]),
    "Redux":("Frontend",["web"]),
    "GraphQL":("Backend",["web"]),
    "REST APIs":("Backend",["tech"]),
    "WebSockets":("Backend",["tech"]),
    "FastAPI":("Backend",["tech","ml"]),
    "Flask":("Backend",["tech","ml"]),
    "Django":("Backend",["tech"]),
    "Spring Boot":("Backend",["tech"]),
    "Express":("Backend",["web"]),
    "JWT":("Security",["tech"]),
    "OAuth":("Security",["tech"]),
    "OAuth2":("Security",["tech"]),
    "Auth0":("Security",["tech"]),
    "Microservices":("Backend",["tech"]),
    "gRPC":("Backend",["tech"]),
    "PyTorch":("ML Frameworks",["ml"]),
    "TensorFlow":("ML Frameworks",["ml"]),
    "Keras":("ML Frameworks",["ml"]),
    "Scikit-learn":("ML Libraries",["ml","data"]),
    "XGBoost":("ML Libraries",["ml","data"]),
    "LightGBM":("ML Libraries",["ml","data"]),
    "CatBoost":("ML Libraries",["ml","data"]),
    "Hugging Face":("NLP",["ml","nlp"]),
    "BERT":("NLP",["ml","nlp"]),
    "GPT":("NLP",["ml","nlp"]),
    "LLM":("NLP",["ml","nlp"]),
    "NLP":("NLP",["ml","nlp"]),
    "Computer Vision":("ML",["ml"]),
    "OpenCV":("ML",["ml"]),
    "LangChain":("AI Tools",["ml","nlp"]),
    "LlamaIndex":("AI Tools",["ml","nlp"]),
    "Machine Learning":("ML",["ml"]),
    "Deep Learning":("ML",["ml"]),
    "Reinforcement Learning":("ML",["ml"]),
    "Transfer Learning":("ML",["ml"]),
    "Feature Engineering":("ML",["ml","data"]),
    "Model Deployment":("MLOps",["ml"]),
    "Pandas":("Data Libraries",["data","ml"]),
    "NumPy":("Data Libraries",["data","ml"]),
    "Matplotlib":("Visualization",["data"]),
    "Seaborn":("Visualization",["data"]),
    "Plotly":("Visualization",["data"]),
    "SciPy":("Data Libraries",["data"]),
    "Kafka":("Data Streaming",["data","tech"]),
    "Spark":("Big Data",["data"]),
    "PySpark":("Big Data",["data"]),
    "Dask":("Big Data",["data","ml"]),
    "Flink":("Data Streaming",["data"]),
    "Airflow":("MLOps",["data","ml"]),
    "dbt":("Data Engineering",["data"]),
    "Prefect":("Data Engineering",["data"]),
    "Hadoop":("Big Data",["data"]),
    "Databricks":("Data Engineering",["data","ml"]),
    "Delta Lake":("Data Engineering",["data"]),
    "PostgreSQL":("Database",["tech","data"]),
    "MySQL":("Database",["tech"]),
    "SQLite":("Database",["tech"]),
    "MongoDB":("Database",["tech"]),
    "Redis":("Database",["tech"]),
    "Cassandra":("Database",["tech"]),
    "DynamoDB":("Database",["tech","cloud"]),
    "Elasticsearch":("Database",["tech"]),
    "Neo4j":("Database",["tech"]),
    "Firestore":("Database",["tech"]),
    "Snowflake":("Data Warehouse",["data","finance"]),
    "BigQuery":("Data Warehouse",["data"]),
    "Redshift":("Data Warehouse",["data"]),
    "AWS":("Cloud",["tech","ml"]),
    "GCP":("Cloud",["tech"]),
    "Azure":("Cloud",["tech"]),
    "EC2":("Cloud",["tech"]),
    "S3":("Cloud",["tech"]),
    "Lambda":("Cloud",["tech"]),
    "SageMaker":("Cloud",["ml"]),
    "Docker":("DevOps",["tech","ml"]),
    "Kubernetes":("DevOps",["tech","ml"]),
    "Terraform":("DevOps",["tech"]),
    "Ansible":("DevOps",["tech"]),
    "Helm":("DevOps",["tech"]),
    "Jenkins":("DevOps",["tech"]),
    "GitHub Actions":("DevOps",["tech"]),
    "CircleCI":("DevOps",["tech"]),
    "ArgoCD":("DevOps",["tech"]),
    "CI/CD":("DevOps",["tech"]),
    "Nginx":("DevOps",["tech"]),
    "MLflow":("MLOps",["ml"]),
    "Kubeflow":("MLOps",["ml"]),
    "DVC":("MLOps",["ml"]),
    "Feast":("MLOps",["ml"]),
    "Tecton":("MLOps",["ml"]),
    "Pinecone":("Vector DB",["ml"]),
    "Weaviate":("Vector DB",["ml"]),
    "Chroma":("Vector DB",["ml"]),
    "FAISS":("Vector DB",["ml"]),
    "Git":("Tools",["tech"]),
    "GitHub":("Tools",["tech"]),
    "GitLab":("Tools",["tech"]),
    "Jira":("Tools",["tech"]),
    "Confluence":("Tools",["tech"]),
    "Notion":("Tools",["tech"]),
    "Postman":("Tools",["tech"]),
    "Jupyter":("Tools",["data","ml"]),
    "Linux":("Tools",["tech"]),
    "Tableau":("Visualization",["data","finance","marketing"]),
    "Power BI":("Visualization",["data","finance"]),
    "Looker":("Visualization",["data","marketing"]),
    "Metabase":("Visualization",["data"]),
    "Grafana":("Monitoring",["tech"]),
    "Datadog":("Monitoring",["tech"]),
    "Prometheus":("Monitoring",["tech"]),
    "Pytest":("Testing",["tech"]),
    "Jest":("Testing",["web"]),
    "Selenium":("Testing",["web"]),
    "Cypress":("Testing",["web"]),
    "DCF":("Finance",["finance"]),
    "LBO":("Finance",["finance"]),
    "Financial Modelling":("Finance",["finance"]),
    "Valuation":("Finance",["finance"]),
    "Due Diligence":("Finance",["finance"]),
    "Bloomberg Terminal":("Finance",["finance"]),
    "Bloomberg":("Finance",["finance"]),
    "Capital IQ":("Finance",["finance"]),
    "FactSet":("Finance",["finance"]),
    "Comparable Company Analysis":("Finance",["finance"]),
    "Precedent Transactions":("Finance",["finance"]),
    "Pitch Books":("Finance",["finance"]),
    "Credit Analysis":("Finance",["finance"]),
    "Debt Structuring":("Finance",["finance"]),
    "Portfolio Management":("Finance",["finance"]),
    "Risk Management":("Finance",["finance"]),
    "Financial Statement Analysis":("Finance",["finance"]),
    "IFRS":("Accounting",["finance"]),
    "Ind AS":("Accounting",["finance"]),
    "GAAP":("Accounting",["finance"]),
    "Private Equity":("Finance",["finance"]),
    "Equity Research":("Finance",["finance"]),
    "ESG":("Finance",["finance"]),
    "CFA":("Certification",["finance"]),
    "FRM":("Certification",["finance"]),
    "Excel":("Tools",["finance","data"]),
    "PowerPoint":("Tools",["finance"]),
    "SEO":("Marketing",["marketing"]),
    "SEM":("Marketing",["marketing"]),
    "Google Ads":("Marketing",["marketing"]),
    "Meta Ads":("Marketing",["marketing"]),
    "Facebook Ads":("Marketing",["marketing"]),
    "Email Marketing":("Marketing",["marketing"]),
    "Content Strategy":("Marketing",["marketing"]),
    "Social Media Marketing":("Marketing",["marketing"]),
    "Performance Marketing":("Marketing",["marketing"]),
    "Growth Marketing":("Marketing",["marketing"]),
    "A/B Testing":("Analytics",["marketing","data"]),
    "Google Analytics":("Analytics",["marketing","data"]),
    "Google Analytics 4":("Analytics",["marketing"]),
    "Mixpanel":("Analytics",["marketing","data"]),
    "Amplitude":("Analytics",["marketing"]),
    "HubSpot":("CRM",["marketing"]),
    "Salesforce":("CRM",["marketing","finance"]),
    "Marketo":("CRM",["marketing"]),
    "Braze":("CRM",["marketing"]),
    "Clevertap":("CRM",["marketing"]),
    "Mailchimp":("CRM",["marketing"]),
    "SEMrush":("Marketing Tools",["marketing"]),
    "Ahrefs":("Marketing Tools",["marketing"]),
    "AppsFlyer":("Marketing Tools",["marketing"]),
    "Adjust":("Marketing Tools",["marketing"]),
    "App Store Optimisation":("Marketing",["marketing"]),
    "ASO":("Marketing",["marketing"]),
    "Attribution Modelling":("Analytics",["marketing"]),
    "Cohort Analysis":("Analytics",["marketing","data"]),
    "Lifecycle Marketing":("Marketing",["marketing"]),
    "Programmatic Advertising":("Marketing",["marketing"]),
    "Segment":("Data Tools",["marketing","data"]),
    "mParticle":("Data Tools",["marketing"]),
    "Optimizely":("Marketing Tools",["marketing"]),
    "Figma":("Design",["design"]),
    "Sketch":("Design",["design"]),
    "Adobe XD":("Design",["design"]),
    "Canva":("Design",["design","marketing"]),
    "Wireframing":("Design",["design"]),
    "Prototyping":("Design",["design"]),
    "User Research":("Product",["product"]),
    "Product Management":("Product",["product"]),
    "Agile":("Practices",["tech"]),
    "Scrum":("Practices",["tech"]),
    "Kanban":("Practices",["tech"]),
    "System Design":("Practices",["tech"]),
}

_SKILL_INDEX = {k.lower(): k for k in SKILL_TAXONOMY}

_SHORT_ALLOWLIST = {
    "sql","git","aws","gcp","css","html","php","r","c","go","vba","jwt",
    "dvc","nlp","llm","aso","esg","cfa","frm","m&a","dbt","api",
}

_NOISE = {
    "a","b","c","d","e","f","g","h","i","j","k","l","m",
    "n","o","p","q","r","s","t","u","v","w","x","y","z",
    "growth","scale","impact","drive","lead","manage","build",
    "strong","deep","broad","large","good","great","solid",
    "data","analysis","management","strategy","process","service",
    "quality","performance","operations","support","development",
    "communication","research","planning","monitoring","reporting",
    "training","review","testing","design",
    "analyst","manager","engineer","developer","designer",
    "senior","junior","head","director","associate","intern","consultant",
    "mba","phd","ca","degree","bachelor","master","preferred",
    "rds","ecs","eks","iam","cdn","b2b","b2c","saas","fmcg","sme",
    "roi","kpi","okr","dsp","cdp",
    "google","meta","amazon","microsoft","apple",
    "zepto","nykaa","swiggy","flipkart","meesho","kedaara",
    "icici","hdfc","india","indian","sector","platform","enterprise",
}

_onet_index: dict[str, str] = {}


@app.on_event("startup")
async def startup():
    global _onet_index
    _onet_index.update(_SKILL_INDEX)
    try:
        import pandas as pd
        xlsx = Path("data/Skills.xlsx")
        csv  = Path("data/onet_skills.csv")
        if xlsx.exists():
            df  = pd.read_excel(xlsx)
            col = next((c for c in df.columns if "element" in c.lower()
                        and "name" in c.lower()), df.columns[3])
        elif csv.exists():
            df, col = pd.read_csv(csv), "Element Name"
        else:
            df = None
        if df is not None:
            for t in df[col].dropna().str.strip().unique():
                if isinstance(t, str) and 2 < len(t) < 80:
                    lc = t.lower()
                    if lc not in _onet_index and lc not in _NOISE:
                        _onet_index[lc] = t
        print(f"[STARTUP] Index ready: {len(_onet_index)} terms")
    except Exception as e:
        print(f"[STARTUP] O*NET failed: {e}")


def extract_text(b64: str, fname: str) -> str:
    if "," in b64:
        b64 = b64.split(",", 1)[1]
    try:
        raw = base64.b64decode(b64)
    except Exception as e:
        print(f"[ERROR] base64: {e}")
        return ""
    f = fname.lower()
    if f.endswith(".pdf"):
        try:
            import fitz
            doc  = fitz.open(stream=raw, filetype="pdf")
            text = "\n".join(p.get_text() for p in doc)
            if text.strip():
                print(f"[INFO] PDF: {len(text)} chars")
                return text
            print("[WARN] PDF empty — may be scanned")
        except Exception as e:
            print(f"[WARN] PDF: {e}")
    if f.endswith(".docx"):
        try:
            import docx
            doc   = docx.Document(io.BytesIO(raw))
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            for tbl in doc.tables:
                for row in tbl.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            parts.append(cell.text.strip())
            text = "\n".join(parts)
            print(f"[INFO] DOCX: {len(text)} chars")
            return text
        except ImportError:
            print("[WARN] pip install python-docx")
        except Exception as e:
            print(f"[WARN] DOCX: {e}")
    try:
        text = raw.decode("utf-8", errors="ignore")
        if text.count(" ") > 20:
            return text
    except Exception:
        pass
    return ""


def _valid_short(term: str, text: str) -> bool:
    if term.lower() not in _SHORT_ALLOWLIST:
        return False
    pat = r"(?<![a-zA-Z0-9])" + re.escape(term) + r"(?![a-zA-Z0-9])"
    if not re.search(pat, text, re.IGNORECASE):
        return False
    if len(term) <= 2:
        ctx = r"(skills|languages|tools|stack)[^\n]{0,200}" + re.escape(term.lower())
        sep = r"[,/|•]\s*" + re.escape(term.lower()) + r"\s*[,/|•\n]"
        if not (re.search(ctx, text.lower()) or re.search(sep, text.lower())):
            return False
    return True


def extract_resume_skills(text: str, jd_skills: list | None = None) -> list:
    found, seen = [], set()

    def add(t: str):
        canon = _SKILL_INDEX.get(t.lower(), t)
        if canon.lower() not in seen and canon.lower() not in _NOISE:
            found.append(canon)
            seen.add(canon.lower())

    tokens = re.split(r"\s+", text.strip())
    for i in range(len(tokens)):
        for size in (3, 2, 1):
            chunk = tokens[i:i+size]
            if len(chunk) < size:
                break
            raw_ng = " ".join(chunk)
            clean  = re.sub(r"^[^\w+#./&-]+|[^\w+#./&-]+$", "", raw_ng).lower()
            if not clean or clean in seen or clean in _NOISE:
                continue
            if clean in _onet_index:
                orig = _onet_index[clean]
                if len(clean) <= 4 and not _valid_short(clean, text):
                    continue
                add(orig)
                break

    if jd_skills:
        tl = text.lower()
        for s in jd_skills:
            if s.lower() in seen:
                continue
            pat = r"(?<![a-zA-Z])" + re.escape(s.lower()) + r"(?![a-zA-Z])"
            if re.search(pat, tl):
                add(s)

    print(f"[INFO] Resume skills: {len(found)}")
    return found


def infer_level(skill: str, text: str) -> str:
    count = len(re.findall(re.escape(skill), text, re.IGNORECASE))
    idx   = text.lower().find(skill.lower())
    ctx   = text[max(0, idx-400):idx+400].lower() if idx >= 0 else ""

    if count >= 1 and any(w in ctx for w in [
        "senior","lead","expert","head of","principal","architect",
        "5+ year","6+ year","7+ year","8+ year",
    ]):
        return "advanced"
    if count >= 1 and any(w in ctx for w in [
        "proficient","3+ year","4+ year","strong","extensive",
    ]):
        return "intermediate"
    if count >= 4: return "advanced"
    if count >= 2: return "intermediate"

    yrs = [int(y) for y in re.findall(r"(\d+)\s*\+?\s*year", text.lower()) if int(y) < 30]
    my  = max(yrs, default=0)
    if my >= 4 and count >= 1: return "advanced"
    if my >= 2 and count >= 1: return "intermediate"

    m = re.search(
        r"(skills|competencies|expertise|technologies|tools)[^\n]{0,50}\n([\s\S]{0,1000})",
        text.lower()
    )
    if m and skill.lower() in m.group(2):
        return "intermediate"
    return "beginner"


def compute_skill_gaps(resume_skills: list, jd_skills: list) -> list:
    from gap_engine import compute_gaps as _cg
    resume_lower = {s.lower() for s in resume_skills}
    true_gaps    = [s for s in jd_skills if s.lower() not in resume_lower]
    if not true_gaps:
        return []
    raw = _cg(resume_skills, true_gaps)
    if not raw:
        return []
    scores = sorted(g["score"] for g in raw)
    n      = len(scores)
    t_high = scores[max(0, n // 3)]
    t_med  = scores[max(0, 2 * n // 3)]

    def pri(s): return "high" if s <= t_high else "medium" if s <= t_med else "low"
    def lv(s):  return "intermediate" if s >= 0.40 else "beginner"

    return [{
        "skill":    g["skill"],
        "score":    g["score"],
        "priority": pri(g["score"]),
        "level":    lv(g["score"]),
        "category": SKILL_TAXONOMY.get(g["skill"], ("Technical", []))[0],
    } for g in raw]


SKILL_DEPS = {
    "Python":("foundation",[]),"SQL":("foundation",[]),"Git":("foundation",[]),
    "Linux":("foundation",[]),"Statistics":("foundation",[]),"Excel":("foundation",[]),
    "Financial Modelling":("foundation",[]),
    "PyTorch":("core",["Python"]),"TensorFlow":("core",["Python"]),
    "Scikit-learn":("core",["Python"]),"Hugging Face":("core",["Python","PyTorch"]),
    "BERT":("core",["Hugging Face"]),"NLP":("core",["Python"]),
    "Feature Engineering":("core",["Python"]),"MLflow":("core",["Python"]),
    "Docker":("core",["Linux"]),"AWS":("core",[]),
    "Snowflake":("core",["SQL"]),"BigQuery":("core",["SQL"]),
    "DCF":("core",["Excel","Financial Modelling"]),
    "LBO":("core",["Excel","Financial Modelling"]),
    "Credit Analysis":("core",["Financial Modelling"]),
    "Bloomberg Terminal":("core",[]),"Capital IQ":("core",[]),
    "Google Ads":("core",[]),"Meta Ads":("core",[]),
    "Google Analytics":("core",[]),"HubSpot":("core",[]),
    "Salesforce":("core",[]),"SEO":("core",[]),
    "Kubernetes":("advanced",["Docker"]),"Airflow":("advanced",["Python"]),
    "Kafka":("advanced",["Python"]),"Spark":("advanced",["Python"]),
    "CI/CD":("advanced",["Git","Docker"]),"Feast":("advanced",["Python"]),
    "Terraform":("advanced",["Docker"]),"SageMaker":("advanced",["AWS","PyTorch"]),
    "Model Deployment":("advanced",["Docker","PyTorch"]),
    "Debt Structuring":("advanced",["DCF","Credit Analysis"]),
    "Comparable Company Analysis":("advanced",["Financial Modelling"]),
    "ESG":("advanced",["Financial Modelling"]),
    "Attribution Modelling":("advanced",["Google Analytics"]),
    "Cohort Analysis":("advanced",["Google Analytics","SQL"]),
    "AppsFlyer":("advanced",["Google Analytics"]),
    "App Store Optimisation":("advanced",[]),
    "Pinecone":("optional",["Python"]),"Weaviate":("optional",["Python"]),
    "LangChain":("optional",["Python","Hugging Face"]),
}

DURATIONS = {
    "PyTorch":"3 weeks","TensorFlow":"3 weeks","Hugging Face":"2 weeks","BERT":"2 weeks",
    "NLP":"3 weeks","MLflow":"1 week","Docker":"1 week","Kubernetes":"2 weeks",
    "Airflow":"2 weeks","Kafka":"2 weeks","Spark":"2 weeks","CI/CD":"1 week",
    "Feast":"1 week","Snowflake":"1 week","BigQuery":"1 week","Terraform":"2 weeks",
    "AWS":"2 weeks","SageMaker":"2 weeks","Feature Engineering":"2 weeks",
    "Model Deployment":"2 weeks","DCF":"2 weeks","LBO":"3 weeks",
    "Financial Modelling":"3 weeks","Credit Analysis":"2 weeks",
    "Bloomberg Terminal":"1 week","Capital IQ":"1 week",
    "Comparable Company Analysis":"2 weeks","Debt Structuring":"2 weeks",
    "ESG":"1 week","Google Ads":"1 week","Meta Ads":"1 week","SEO":"2 weeks",
    "HubSpot":"1 week","Salesforce":"1 week","Google Analytics":"1 week",
    "Attribution Modelling":"2 weeks","Cohort Analysis":"1 week",
    "AppsFlyer":"1 week","App Store Optimisation":"2 weeks",
}

DESCRIPTIONS = {
    "PyTorch":"Tensors, autograd, and neural network fundamentals using PyTorch.",
    "TensorFlow":"Build and train deep learning models with TensorFlow and Keras.",
    "Hugging Face":"Pretrained models for NLP tasks with the Transformers library.",
    "BERT":"Fine-tune BERT for classification, NER, and Q&A tasks.",
    "NLP":"Tokenization, embeddings, sequence models, and text pipelines.",
    "MLflow":"Experiment tracking, model registry, and ML lifecycle management.",
    "Docker":"Containerise apps for reproducible, portable deployments.",
    "Kubernetes":"Orchestrate containers at scale for production services.",
    "Airflow":"Schedule and monitor data and ML pipelines as DAGs.",
    "Kafka":"Build real-time event-driven data pipelines with Kafka.",
    "Spark":"Distributed data processing and large-scale analytics with PySpark.",
    "CI/CD":"Automate build, test, and deploy pipelines for faster delivery.",
    "Feast":"Feature store — manage, serve, and share ML features.",
    "Snowflake":"Cloud data warehousing and SQL analytics at scale.",
    "BigQuery":"Large-scale analytics on Google Cloud with BigQuery SQL.",
    "Terraform":"Infrastructure-as-code for repeatable cloud environments.",
    "AWS":"Core services: compute, storage, networking, and ML on AWS.",
    "SageMaker":"End-to-end ML on AWS: training, tuning, and deployment.",
    "Feature Engineering":"Transform raw data into features that improve model performance.",
    "Model Deployment":"Serve ML models in production via APIs and cloud platforms.",
    "DCF":"Discounted Cash Flow modelling for company and asset valuation.",
    "LBO":"Leveraged Buyout modelling for private equity transactions.",
    "Financial Modelling":"Build robust three-statement and scenario models in Excel.",
    "Credit Analysis":"Assess borrower risk through cash flow and balance sheet analysis.",
    "Bloomberg Terminal":"Navigate Bloomberg for market data, news, and analytics.",
    "Capital IQ":"Use S&P Capital IQ for company data, comps, and deals.",
    "Comparable Company Analysis":"Value companies using peer multiples and benchmarks.",
    "Debt Structuring":"Design and analyse debt instruments and capital structures.",
    "ESG":"Integrate ESG factors in investment decisions and portfolio reporting.",
    "Google Ads":"Campaign setup, bidding strategy, and performance optimisation.",
    "Meta Ads":"Audience targeting and creative optimisation on Facebook/Instagram.",
    "SEO":"Technical SEO, keyword strategy, and on-page optimisation.",
    "HubSpot":"CRM workflows, email sequences, and marketing automation.",
    "Salesforce":"CRM configuration, lead management, and sales pipeline tracking.",
    "Google Analytics":"Funnel analysis, event tracking, and audience insights on GA4.",
    "Attribution Modelling":"Multi-touch attribution to accurately measure channel ROI.",
    "Cohort Analysis":"Track retention and behaviour patterns across user cohorts.",
    "AppsFlyer":"Mobile attribution, deep linking, and campaign measurement.",
    "App Store Optimisation":"Keyword optimisation and creative testing for app stores.",
}


def build_roadmap(gap_skills: list) -> list:
    if not gap_skills:
        return []
    try:
        courses = json.loads(Path("data/courses.json").read_text())
        clookup = {}
        for c in courses:
            for s in c.get("skills_covered", []):
                clookup[s.lower()] = c
    except Exception:
        clookup = {}

    expanded: set[str] = set()

    def canon(n: str) -> str:
        return _SKILL_INDEX.get(n.lower(), n)

    def add_deps(sk: str):
        if sk in expanded:
            return
        expanded.add(sk)
        _, deps = SKILL_DEPS.get(sk, ("core", []))
        for d in deps:
            add_deps(d)

    for g in gap_skills:
        add_deps(canon(g["name"]))
    if not expanded:
        return []

    s2id  = {s: f"n{i+1}" for i, s in enumerate(sorted(expanded))}
    nodes = []
    for sk in expanded:
        nid      = s2id[sk]
        st, deps = SKILL_DEPS.get(sk, ("core", []))
        c        = clookup.get(sk.lower())
        if not c:
            for ck, cv in clookup.items():
                if sk.lower() in ck or ck in sk.lower():
                    c = cv
                    break
        children = [s2id[s] for s in expanded
                    if sk in SKILL_DEPS.get(s, ("core", []))[1]]
        nodes.append({
            "id":          nid,
            "title":       c["title"] if c else f"Learn {sk}",
            "description": DESCRIPTIONS.get(sk, f"Build practical skills in {sk}."),
            "duration":    DURATIONS.get(sk, "1 week"),
            "type":        st,
            "status":      "locked" if any(d in expanded for d in deps) else "available",
            "children":    children,
        })

    order = {"foundation": 0, "core": 1, "advanced": 2, "optional": 3}
    nodes.sort(key=lambda n: order.get(n["type"], 2))
    return nodes


def _jd_skills(jd: str) -> list:
    from gap_engine import extract_skills_from_jd
    return extract_skills_from_jd(jd)

def _target_role(jd: str) -> str:
    for line in jd.strip().split("\n")[:8]:
        line = line.strip()
        if any(k in line.lower() for k in ["title:", "role:", "position:"]):
            p = line.split(":", 1)
            if len(p) > 1 and p[1].strip():
                return p[1].strip()
        if 8 < len(line) < 60 and re.match(r"^[A-Z]", line):
            return line
    return "Target Role"

def _candidate_name(text: str) -> str:
    for line in text.strip().split("\n")[:5]:
        line = line.strip()
        if 5 < len(line) < 50 and re.match(r"^[A-Z][a-z]+ [A-Z]", line):
            return line
    return "Candidate"


@app.get("/health")
def health():
    return {"status":"ok",
            "model_ready":Path("models/skill_extractor/config.json").exists(),
            "index_size":len(_onet_index)}


@app.post("/analyze")
async def analyze(
    resume_b64:  str = Form(...),
    resume_name: str = Form(...),
    jd_text:     str = Form(...),
):
    try:
        print(f"\n[/analyze] {resume_name} | jd={len(jd_text)} chars")
        text = extract_text(resume_b64, resume_name)
        if not text.strip():
            return {"error": "Could not read resume. Use a text-based PDF or DOCX."}

        jd_sk  = _jd_skills(jd_text)
        print(f"[INFO] JD: {len(jd_sk)} skills → {jd_sk[:5]}")
        res_sk = extract_resume_skills(text, jd_skills=jd_sk)
        gaps   = compute_skill_gaps(res_sk, jd_sk)

        gap_set = {g["skill"] for g in gaps}
        matched = [{"name":s, "level":infer_level(s,text),
                    "category":SKILL_TAXONOMY.get(s,("Technical",[]))[0]}
                   for s in res_sk if s not in gap_set]
        gap_out = [{"name":g["skill"],"level":g["level"],
                    "category":g["category"],"priority":g["priority"]}
                   for g in gaps]

        total = len(matched) + len(gap_out)
        score = round(len(matched) / total * 100) if total > 0 else 0
        print(f"[INFO] matched={len(matched)} gaps={len(gap_out)} score={score}%")

        return {"candidateName":_candidate_name(text),
                "targetRole":_target_role(jd_text),
                "matchScore":score,
                "matchedSkills":matched,
                "gapSkills":gap_out}

    except Exception as e:
        traceback.print_exc()
        return {"error": str(e)}


class RoadmapRequest(BaseModel):
    gapSkills: list[dict]


@app.post("/roadmap")
def roadmap(body: RoadmapRequest):
    try:
        nodes = build_roadmap(body.gapSkills)
        print(f"[/roadmap] {len(body.gapSkills)} gaps → {len(nodes)} nodes")
        return nodes
    except Exception as e:
        traceback.print_exc()
        return []


if __name__ == "__main__":
    import uvicorn
    uvicorn.run("api:app", host="0.0.0.0", port=8000, reload=True)
